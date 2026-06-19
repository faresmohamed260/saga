from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt


OUTPUT_PATH = Path(r"B:\Downloads\Updated_Section_2_7_Tools_Research.docx")


SECTION_CONTENT: list[tuple[str, str]] = [
    (
        "2.7 Tools Research",
        "",
    ),
    (
        "2.7.1 Large Language Models",
        "The current S.A.G.A. system does not depend on a single fixed language model. Instead, it uses a provider-agnostic runtime layer that allows different model backends to be selected for different tasks such as narrative analysis, planning, validation, and prose generation.",
    ),
    (
        "Mistral and Other Open-Weight Models",
        "Open-weight models such as Mistral remain useful because they offer lower operational cost, self-hosting flexibility, and strong performance for structured extraction and experimentation. These models are especially valuable in local or semi-local workflows where data locality and cost control matter. Their main limitations are lower peak reasoning quality than stronger hosted models and greater sensitivity to prompt and validation design.",
    ),
    (
        "OpenAI-Compatible Hosted Models",
        "Hosted models provide strong reasoning, reliable schema-following behavior, and mature API ergonomics. This makes them well suited for planning-heavy tasks, constrained generation, and validation loops where output correctness is especially important. Their main limitations are higher cost, dependence on external APIs, and reduced control over runtime internals compared with fully local deployments.",
    ),
    (
        "Meta LLaMA and Similar Open-Weight Families",
        "Meta LLaMA and related model families provide another strong self-hosted option with broad ecosystem support. They are suitable for experimentation, retrieval-assisted workflows, and selective prose generation, but they usually require more hardware resources and operational tuning than lighter open-weight alternatives.",
    ),
    (
        "Current Selection Rationale",
        "In the current system, model choice is abstracted through the LLMClient layer instead of being hardwired into the architecture. This allows planning, prose generation, and analysis workloads to use different providers as reliability, cost, and quality requirements evolve. This provider-agnostic design better reflects the production path of S.A.G.A. than the earlier prototype's fixed-model framing.",
    ),
    (
        "2.7.2 Identity Analysis Tool",
        "BookNLP is a core tool in the current S.A.G.A. architecture because it serves as the foundation of the production identity pipeline. The system uses a BookNLP-clean identity workflow in which raw BookNLP output is adapted, cleaned, and normalized before deeper narrative analysis begins. This process helps establish stable canonical characters, alias mappings, narrator handling, and reference entities before events, relationships, profiles, and states are extracted by later stages.",
    ),
    (
        "Why BookNLP Was Selected",
        "BookNLP was selected because the current system requires identity stability before downstream analysis can be trusted. Relying only on scene-local extraction caused fragmentation and inconsistency in the older prototype. By moving identity upstream and making BookNLP-clean the supported production path, S.A.G.A. improves character continuity, reduces duplicate identities, and provides a stable basis for retrieval, visual continuity, and decoder generation.",
    ),
    (
        "2.7.3 Database",
        "SQLite is selected as the canonical operational database because it provides a simple, durable, and queryable local store that fits the current architecture of S.A.G.A. Unlike the earlier graph-first prototype, the current system persists books, chapters, scenes, identities, entities, events, relationships, timelines, stable states, visual prompts, generated stories, and runtime jobs directly into a normalized SQLite schema. This reduces synchronization overhead, simplifies deployment, improves resumability, and gives the dashboard and downstream services one authoritative source of truth. Optional graph-oriented services may still be used for selected retrieval experiments, but SQLite is the primary persistence layer of the current system.",
    ),
    (
        "2.7.4 Orchestration and Runtime Tools",
        "The current orchestration stack is implemented through a layered local runtime rather than external workflow tools. FastAPI provides the HTTP runtime and job-control surface for uploads, analysis, decoder workflows, visual workflows, and diagnostics. Dashboard Pro, implemented in React, provides the operator-facing interface for staging uploads, monitoring runs, inspecting analysis outputs, reviewing generated stories, browsing visual assets, and checking provider health. DatabaseAnalysisRunService coordinates validation, ingestion, identity preparation, stage execution, and progress tracking for the database-native pipeline. Together, these components form the operational backbone of S.A.G.A.",
    ),
    (
        "2.7.5 Visual Generation Tool",
        "ComfyUI is used in the current system as the main downstream visual-generation environment. It is not part of the canonical analysis core itself; instead, it receives structured prompt payloads generated from the persisted canonical store. Through services such as ComfyUICharacterSheetService and the visual-assets workflow in the dashboard, S.A.G.A. can transform canonical character baselines, scene-aware state, and stored prompt versions into render-ready visual outputs.",
    ),
    (
        "Why ComfyUI Was Selected",
        "ComfyUI was selected because it offers flexible workflow-based image generation while remaining compatible with structured prompt construction and repeated visual iteration. This makes it suitable for character sheets, scene imagery, and other grounded visual assets that depend on continuity rather than one-shot prompt generation. In the current architecture, ComfyUI is best understood as a rendering layer that operates on top of S.A.G.A.'s canonical memory.",
    ),
    (
        "2.7.6 Cloud Rendering Infrastructure",
        "Modal is used as part of the ComfyUI integration layer rather than as the main system platform. In the current project, Modal provides cloud-hosted GPU execution for ComfyUI workloads, together with token-rotation and failover support for team-based rendering. This allows visual-generation jobs to be dispatched to remote GPU-backed environments while keeping the main narrative-analysis and dashboard runtime local.",
    ),
    (
        "Why Modal Was Selected",
        "Modal was selected because image-generation workloads require GPU resources that may not always be practical to dedicate locally. By combining Modal with ComfyUI, the project can run render workflows on remote GPU infrastructure while maintaining a controlled local orchestration layer. This separation keeps the core S.A.G.A. pipeline focused on analysis, persistence, and review, while allowing the rendering subsystem to scale independently when visual assets are needed.",
    ),
    (
        "2.7.7 Source Processing and Ingestion Tools",
        "Among the evaluated source-processing approaches, the current system relies on a pragmatic file-type-specific ingestion layer rather than a vision-first OCR pipeline. EPUB sources are handled through EbookLib-based extraction, while PDF-oriented inputs are processed using PyMuPDF and supporting PDF utilities such as pypdf. Both source paths feed a shared chapter and scene preparation pipeline that prioritizes durable ingestion, provenance preservation, and compatibility with later narrative analysis stages.",
    ),
]


def set_base_style(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)


def add_title(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = paragraph.add_run(text)
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(16)


def add_heading(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)


def add_body(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(text)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY


def main() -> None:
    document = Document()
    set_base_style(document)

    title, _ = SECTION_CONTENT[0]
    add_title(document, title)

    for heading, body in SECTION_CONTENT[1:]:
        add_heading(document, heading)
        add_body(document, body)

    document.save(OUTPUT_PATH)
    print(OUTPUT_PATH)


if __name__ == "__main__":
    main()
